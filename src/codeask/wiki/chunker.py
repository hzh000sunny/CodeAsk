"""Parse documents into normalized wiki chunks."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from markdown_it import MarkdownIt

from codeask.wiki.signals import extract_signals
from codeask.wiki.tokenizer import to_ngrams, tokenize

MAX_CHUNK_CHARS = 3000


@dataclass(slots=True)
class ParsedChunk:
    chunk_index: int
    heading_path: str
    raw_text: str
    normalized_text: str
    tokenized_text: str
    ngram_text: str
    signals_json: dict[str, list[str]]
    start_offset: int
    end_offset: int


def _build(index: int, heading_path: str, raw: str, start: int, end: int) -> ParsedChunk:
    normalized = " ".join(raw.split())
    return ParsedChunk(
        chunk_index=index,
        heading_path=heading_path,
        raw_text=raw,
        normalized_text=normalized,
        tokenized_text=tokenize(normalized),
        ngram_text=to_ngrams(normalized),
        signals_json=extract_signals(raw),
        start_offset=start,
        end_offset=end,
    )


def _split_oversize(text: str) -> list[str]:
    if len(text) <= MAX_CHUNK_CHARS:
        return [text]
    parts: list[str] = []
    buffer = ""
    for paragraph in (part for part in text.split("\n\n") if part.strip()):
        if len(paragraph) > MAX_CHUNK_CHARS:
            if buffer.strip():
                parts.append(buffer.strip())
                buffer = ""
            parts.extend(
                paragraph[index : index + MAX_CHUNK_CHARS]
                for index in range(0, len(paragraph), MAX_CHUNK_CHARS)
            )
            continue
        if len(buffer) + len(paragraph) + 2 > MAX_CHUNK_CHARS and buffer:
            parts.append(buffer.strip())
            buffer = paragraph
        else:
            buffer = f"{buffer}\n\n{paragraph}" if buffer else paragraph
    if buffer.strip():
        parts.append(buffer.strip())
    return parts


class DocumentChunker:
    """Stateless chunker for supported wiki document kinds."""

    def __init__(self) -> None:
        self._md = MarkdownIt("commonmark")

    def chunk_file(self, path: Path, kind: str) -> list[ParsedChunk]:
        if kind == "markdown":
            return self.chunk_markdown(path.read_text(encoding="utf-8"))
        if kind == "text":
            return self.chunk_text(path.read_text(encoding="utf-8"))
        if kind == "pdf":
            return self.chunk_pdf(path)
        if kind == "docx":
            return self.chunk_docx(path)
        raise ValueError(f"unsupported document kind: {kind}")

    def chunk_markdown(self, source: str) -> list[ParsedChunk]:
        tokens = self._md.parse(source)
        source_lines = source.splitlines()
        sections: list[tuple[str, str]] = []
        heading_stack: list[tuple[int, str]] = []
        current_heading_path = ""
        current_body_lines: list[str] = []

        index = 0
        while index < len(tokens):
            token = tokens[index]
            if token.type == "heading_open":
                if current_body_lines:
                    sections.append((current_heading_path, "\n".join(current_body_lines).strip()))
                    current_body_lines = []
                level = int(token.tag[1])
                inline = tokens[index + 1]
                heading_text = inline.content.strip()
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, heading_text))
                current_heading_path = " > ".join(heading for _, heading in heading_stack)
                index += 3
                continue

            if token.type in ("paragraph_open", "bullet_list_open", "ordered_list_open"):
                if token.map is not None:
                    start_line, end_line = token.map
                    snippet = "\n".join(source_lines[start_line:end_line])
                    if snippet.strip():
                        current_body_lines.append(snippet)
                close_type = token.type.replace("_open", "_close")
                depth = 1
                index += 1
                while index < len(tokens) and depth > 0:
                    if tokens[index].type == token.type:
                        depth += 1
                    elif tokens[index].type == close_type:
                        depth -= 1
                    index += 1
                continue

            if token.type in ("fence", "code_block"):
                fence = token.content.rstrip("\n")
                language = token.info.strip() if token.info else ""
                rendered = f"```{language}\n{fence}\n```" if token.type == "fence" else fence
                current_body_lines.append(rendered)
            elif token.type == "hr":
                current_body_lines.append("---")
            index += 1

        if current_body_lines:
            sections.append((current_heading_path, "\n".join(current_body_lines).strip()))

        chunks: list[ParsedChunk] = []
        offset = 0
        chunk_index = 0
        for heading_path, body in sections:
            if not body:
                continue
            for piece in _split_oversize(body):
                start = offset
                end = start + len(piece)
                chunks.append(_build(chunk_index, heading_path, piece, start, end))
                chunk_index += 1
                offset = end + 2
        return chunks

    def chunk_text(self, source: str) -> list[ParsedChunk]:
        chunks: list[ParsedChunk] = []
        offset = 0
        chunk_index = 0
        for paragraph in source.split("\n\n"):
            body = paragraph.strip()
            if not body:
                offset += len(paragraph) + 2
                continue
            for piece in _split_oversize(body):
                start = offset
                end = start + len(piece)
                chunks.append(_build(chunk_index, "", piece, start, end))
                chunk_index += 1
                offset = end + 2
        return chunks

    def chunk_pdf(self, path: Path) -> list[ParsedChunk]:
        import pypdfium2 as pdfium  # pyright: ignore[reportMissingTypeStubs]

        pdf = pdfium.PdfDocument(str(path))
        try:
            full_text_pieces: list[str] = []
            for page in pdf:
                text_page = page.get_textpage()
                try:
                    full_text_pieces.append(text_page.get_text_range())
                finally:
                    text_page.close()
                    page.close()
            return self.chunk_text("\n\n".join(full_text_pieces))
        finally:
            pdf.close()

    def chunk_docx(self, path: Path) -> list[ParsedChunk]:
        import docx

        document = docx.Document(str(path))
        sections: list[tuple[str, list[str]]] = []
        current_heading = ""
        current_paragraphs: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style
            style_name = style.name if style is not None and style.name is not None else ""
            if style_name.startswith("Heading"):
                if current_paragraphs:
                    sections.append((current_heading, current_paragraphs))
                    current_paragraphs = []
                current_heading = text
            else:
                current_paragraphs.append(text)
        if current_paragraphs:
            sections.append((current_heading, current_paragraphs))

        chunks: list[ParsedChunk] = []
        offset = 0
        chunk_index = 0
        for heading, paragraphs in sections:
            body = "\n\n".join(paragraphs)
            for piece in _split_oversize(body):
                start = offset
                end = start + len(piece)
                chunks.append(_build(chunk_index, heading, piece, start, end))
                chunk_index += 1
                offset = end + 2
        return chunks
